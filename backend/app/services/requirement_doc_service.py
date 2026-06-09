"""
需求文档服务
============
负责上传、保存、解析、分类管理需求文档，并为 AI 生成 case 提供上下文。
"""

import hashlib
import json
import os
import uuid
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import UploadFile
from sqlalchemy.sql import func as sql_func
from sqlalchemy.orm import Session

from app.config import settings
from app.models.requirement_document import RequirementDocument
from app.models.requirement_issue import RequirementIssue, RequirementIssueAction, RequirementRevision
from app.models.requirement_item import RequirementItem
from app.models.requirement_tree_node import RequirementTreeNode
from app.models.test_case import TestCase
from app.models.changelog import Changelog
from app.models.api_info import APIInfo
from app.services.requirement_tree_service import RequirementTreeService
from app.skills.api_document_parse_skill import ApiDocumentParseSkill
from app.skills.api_log_parse_skill import ApiLogParseSkill
from app.skills.case_generate_skill import CaseGenerateSkill
from app.skills.document_parse_router_skill import DocumentParseRouterSkill
from app.skills.requirement_diff_skill import RequirementDiffSkill
from app.skills.requirement_parse_skill import RequirementParseSkill


class RequirementBlockingIssuesError(ValueError):
    """Raised when unresolved blocking requirement issues prevent an operation."""

    def __init__(self, message: str, issues: List[Dict[str, Any]]):
        super().__init__(message)
        self.issues = issues

    def to_detail(self) -> Dict[str, Any]:
        return {"message": str(self), "blocking_issues": self.issues}


class RequirementDocService:
    """需求文档服务"""

    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}
    PARSE_STATUSES = {"unparsed", "parsing", "pending_review", "check_failed", "stored"}

    def __init__(self, db: Session):
        self.db = db
        self.upload_dir = Path(settings.DOCUMENT_UPLOAD_DIR)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    async def save_document(
        self,
        file: UploadFile,
        title: str,
        category: str,
        module: str,
        dependency_scope: str,
        dependency_notes: str,
        tags: str,
        created_by: str,
        tree_node_id: Optional[int] = None,
        project_id: Optional[int] = None,
    ) -> RequirementDocument:
        """保存上传文档并提取文本"""
        if tree_node_id is not None:
            node = RequirementTreeService(self.db).get_node(tree_node_id)
            if not node:
                raise ValueError("树节点不存在")
        suffix = Path(file.filename or "").suffix.lower()
        if suffix not in self.ALLOWED_EXTENSIONS:
            raise ValueError("仅支持 pdf、docx、md、txt 文件")

        file_id = uuid.uuid4().hex
        saved_name = f"{file_id}{suffix}"
        saved_path = self.upload_dir / saved_name

        content = await file.read()
        file_hash = hashlib.sha256(content).hexdigest()
        duplicate = (
            self.db.query(RequirementDocument)
            .filter(
                RequirementDocument.file_hash == file_hash,
                RequirementDocument.status == "active",
            )
            .first()
        )
        if duplicate:
            raise ValueError(f"检测到重复文件，已存在文档《{duplicate.title}》")

        doc_title = title or Path(file.filename or saved_name).stem
        title_duplicate = (
            self.db.query(RequirementDocument)
            .filter(
                RequirementDocument.title == doc_title,
                RequirementDocument.status == "active",
            )
            .first()
        )
        if title_duplicate:
            raise ValueError(f"已存在同名文档《{title_duplicate.title}》，请修改标题或归档旧文档")

        with open(saved_path, "wb") as output:
            output.write(content)

        extracted_content = self._extract_text(saved_path, suffix)

        document = RequirementDocument(
            title=doc_title,
            file_name=file.filename or saved_name,
            file_type=suffix.replace(".", ""),
            file_path=str(saved_path),
            file_size=len(content),
            file_hash=file_hash,
            project_id=project_id,
            tree_node_id=tree_node_id,
            category=category or "未分类",
            module=module or "",
            dependency_scope=dependency_scope or "",
            dependency_notes=dependency_notes or "",
            tags=tags or "",
            extracted_content=extracted_content,
            parse_status="unparsed",
            created_by=created_by or "",
        )
        self.db.add(document)
        self.db.commit()
        self.db.refresh(document)
        return document

    def list_documents(
        self,
        category: Optional[str] = None,
        keyword: Optional[str] = None,
        tree_node_id: Optional[int] = None,
        status: str = "active",
        page: int = 1,
        page_size: int = 20,
    ) -> Dict:
        query = self.db.query(RequirementDocument).filter(RequirementDocument.status == status)
        if category:
            query = query.filter(RequirementDocument.category == category)
        if tree_node_id is not None:
            query = query.filter(RequirementDocument.tree_node_id == tree_node_id)
        if keyword:
            query = query.filter(
                (RequirementDocument.title.contains(keyword))
                | (RequirementDocument.file_name.contains(keyword))
                | (RequirementDocument.module.contains(keyword))
            )

        total = query.count()
        items = (
            query.order_by(RequirementDocument.created_at.desc())
            .offset((page - 1) * page_size)
            .limit(page_size)
            .all()
        )
        return {
            "total": total,
            "page": page,
            "page_size": page_size,
            "items": [item.to_dict() for item in items],
        }

    def update_document_meta(
        self,
        doc_id: int,
        title: Optional[str] = None,
        category: Optional[str] = None,
        module: Optional[str] = None,
    ) -> RequirementDocument:
        document = self.get_document(doc_id)
        if not document or document.status != "active":
            raise ValueError("需求文档不存在")
        if title is not None:
            document.title = title
        if category is not None:
            document.category = category
        if module is not None:
            document.module = module
        self.db.commit()
        self.db.refresh(document)
        return document

    def archive_document(self, doc_id: int) -> RequirementDocument:
        document = self.get_document(doc_id)
        if not document or document.status != "active":
            raise ValueError("需求文档不存在")
        document.status = "archived"
        self.db.commit()
        self.db.refresh(document)
        return document

    def soft_delete_document(self, doc_id: int) -> RequirementDocument:
        document = self.get_document(doc_id)
        if not document or document.status != "active":
            raise ValueError("需求文档不存在")
        document.status = "deleted"
        self.db.commit()
        self.db.refresh(document)
        return document

    def restore_document(self, doc_id: int) -> RequirementDocument:
        document = self.db.query(RequirementDocument).filter(RequirementDocument.id == doc_id).first()
        if not document or document.status not in {"archived", "deleted"}:
            raise ValueError("文档不在归档或回收站中")
        document.status = "active"
        self.db.commit()
        self.db.refresh(document)
        return document

    def can_generate_cases(self, doc_id: int) -> bool:
        document = self.get_document(doc_id)
        return bool(
            document
            and document.status == "active"
            and document.parse_status == "stored"
            and not self.has_unresolved_blocking_issues(doc_id)
        )

    def assert_can_generate_cases(self, doc_id: int) -> None:
        document = self.get_document(doc_id)
        if not document:
            raise ValueError("需求文档不存在")
        blocking_issues = self.list_unresolved_blocking_issues(doc_id)
        if blocking_issues:
            raise RequirementBlockingIssuesError("文档存在未解决的阻断级问题，不能生成 case", blocking_issues)
        if document.parse_status != "stored":
            raise ValueError("文档尚未通过解析检查并入库，不能生成 case")

    def trigger_parse(self, doc_id: int) -> Dict[str, Any]:
        document = self.get_document(doc_id)
        if not document or document.status != "active":
            raise ValueError("需求文档不存在")

        document.parse_status = "parsing"
        self.db.commit()
        try:
            route = DocumentParseRouterSkill().run(
                {
                    "content": document.extracted_content or "",
                    "file_name": document.file_name or "",
                    "title": document.title or "",
                    "category": document.category or "未分类",
                }
            ).get("data", {})
            document_type = route.get("document_type", "requirement_document")
            parser_skill = route.get("parser_skill", "RequirementParseSkill")
            parse_payload = self._run_bound_parse_skill(document, document_type, parser_skill)
            points = parse_payload.get("requirement_points", [])
            issues = parse_payload.get("issues", [])

            # 获取该文档下所有已被用户显式确认/忽略，且原文片段仍一致的历史问题项。
            # 不能只按 requirement_no 继承，否则同一编号内容变更后会误压制新的复核问题。
            confirmed_issues = self.db.query(RequirementIssue).filter(
                RequirementIssue.document_id == doc_id,
                RequirementIssue.status.in_(["resolved", "ignored"])
            ).all()
            confirmed_issue_excerpts = {
                (issue.source_location, issue.source_excerpt)
                for issue in confirmed_issues
                if issue.source_location and issue.source_excerpt
            }

            # 如果需求编号和原文片段均已被确认/忽略，则避免重复生成同一问题项。
            for point in points:
                req_no = point.get("requirement_no")
                source_text = point.get("source_text") or point.get("content") or ""
                if req_no and (req_no, source_text) in confirmed_issue_excerpts:
                    point["need_review"] = False
                    point["confirmed"] = True

            if document_type == "requirement_document":
                issues.extend(self._build_parse_issues(points, document))
            elif document_type == "api_document" and not points:
                issues.append({"type": "missing", "severity": "high", "message": "接口文档未解析出接口清单，请检查 Method/Path 表格或接口标题"})
            elif document_type == "api_log" and not parse_payload.get("entries"):
                issues.append({"type": "missing", "severity": "high", "message": "线上接口 log 未解析出请求记录"})
            history_diff = self._build_history_diff(document, points)
            coupling_results = self._build_coupling_results(document, points)
            if history_diff.get("conflict_count", 0):
                for item in history_diff.get("diff_items", []):
                    if item.get("change_type") == "modified":
                        old_node = item.get("old") or {}
                        old_title = old_node.get("title") or old_node.get("content", "")[:15]
                        old_title = str(old_title).lstrip("# ").strip()
                        old_no = old_node.get("requirement_no", "")
                        issues.append(
                            {
                                "type": "conflict",
                                "message": f"与历史需求《{old_title}》 ({old_no}) 存在差异，需人工确认",
                                "requirement_no": old_no,
                                "requirement_title": old_title,
                            }
                        )

            self.db.query(RequirementItem).filter(RequirementItem.document_id == doc_id).delete()
            for point in points:
                self.db.add(
                    RequirementItem(
                        document_id=doc_id,
                        requirement_no=point.get("requirement_no", ""),
                        title=point.get("title", ""),
                        content=point.get("content", ""),
                        source_text=point.get("source_text", ""),
                        priority=point.get("priority", "P1"),
                        item_type=point.get("type", "requirement"),
                        need_review=1 if point.get("need_review") else 0,
                        confirmed=1 if point.get("confirmed") else 0,
                    )
                )
            self.db.flush()
            issues.extend(self._build_need_review_issues(points))
            self._sync_requirement_issues(document, issues)

            document.parse_result = json.dumps(
                {
                    "document_type": document_type,
                    "parser_skill": parser_skill,
                    "parser_confidence": route.get("confidence"),
                    "issues": issues,
                    "requirement_points": points,
                    "skill_output": parse_payload.get("skill_output", {}),
                    "history_diff": history_diff,
                    "coupling_results": coupling_results,
                    "total": len(points),
                },
                ensure_ascii=False,
            )
            document.parse_status = "check_failed" if self.has_unresolved_blocking_issues(doc_id) else "pending_review"
            if self._is_unusable_ai_summary(document.ai_summary):
                document.ai_summary = self._build_rule_parse_summary(points, issues)
            self.db.commit()
            self.db.refresh(document)

            self._sync_issue_todos(document)

            return {
                "document": document.to_dict(),
                "issues": self.list_requirement_issues(doc_id),
                "requirement_points": [item.to_dict() for item in document.requirement_items],
            }
        except ValueError as exc:
            document.parse_status = "check_failed"
            document.parse_result = json.dumps(
                {
                    "document_type": "unknown",
                    "parser_skill": "RequirementParseSkill",
                    "error": str(exc),
                    "issues": [{"type": "runtime_error", "severity": "high", "message": str(exc)}],
                    "requirement_points": [],
                },
                ensure_ascii=False,
            )
            self.db.commit()
            raise

    def confirm_storage(self, doc_id: int) -> RequirementDocument:
        document = self.get_document(doc_id)
        if not document or document.status != "active":
            raise ValueError("需求文档不存在")
        blocking_issues = self.list_unresolved_blocking_issues(doc_id)
        if blocking_issues:
            raise RequirementBlockingIssuesError("文档存在未解决的阻断级问题，不能入库落档", blocking_issues)
        if document.parse_status != "pending_review":
            raise ValueError("仅待确认状态的文档可确认入库")
        for item in document.requirement_items:
            item.confirmed = 1
        document.parse_status = "stored"
        self.db.commit()
        self.db.refresh(document)

        from app.services.todo_service import TodoService
        TodoService(self.db).resolve_todo("REQ_CONFLICT", doc_id, reason="产品经理手动澄清并确认需求，自动核销")

        return document

    def get_document_impact(self, doc_id: int) -> Dict[str, int]:
        document = self.get_document(doc_id)
        if not document:
            raise ValueError("需求文档不存在")
        item_count = self.db.query(RequirementItem).filter(RequirementItem.document_id == doc_id).count()
        case_count = (
            self.db.query(TestCase)
            .filter(TestCase.source_document_id == doc_id, TestCase.is_active == 1)
            .count()
        )
        api_ids = [
            row[0]
            for row in self.db.query(TestCase.api_id)
            .filter(TestCase.source_document_id == doc_id, TestCase.api_id.isnot(None))
            .distinct()
            .all()
            if row[0]
        ]
        change_record_count = self.db.query(Changelog).filter(Changelog.api_id.in_(api_ids)).count() if api_ids else 0
        return {
            "document_id": doc_id,
            "requirement_item_count": item_count,
            "case_count": case_count,
            "knowledge_fragment_count": item_count,
            "change_record_count": change_record_count,
        }

    def get_document_relations(self, doc_id: int) -> Dict[str, Any]:
        document = self.get_document(doc_id)
        if not document:
            raise ValueError("需求文档不存在")
        cases = (
            self.db.query(TestCase)
            .filter(TestCase.source_document_id == doc_id)
            .order_by(TestCase.id.asc())
            .all()
        )
        api_ids = sorted({case.api_id for case in cases if case.api_id})
        apis = self.db.query(APIInfo).filter(APIInfo.id.in_(api_ids)).all() if api_ids else []
        changes = self.db.query(Changelog).filter(Changelog.api_id.in_(api_ids)).all() if api_ids else []
        modules = sorted({api.module for api in apis if api.module})
        return {
            "document": document.to_dict(),
            "cases": [case.to_dict() for case in cases],
            "apis": [api.to_dict() for api in apis],
            "modules": modules,
            "changes": [change.to_dict() for change in changes],
        }

    def list_requirement_issues(self, doc_id: int, status: Optional[str] = None) -> List[Dict[str, Any]]:
        document = self.get_document(doc_id)
        if not document:
            raise ValueError("需求文档不存在")
        query = self.db.query(RequirementIssue).filter(RequirementIssue.document_id == doc_id)
        if status and status != "all":
            query = query.filter(RequirementIssue.status == status)
        elif not status:
            query = query.filter(RequirementIssue.status.in_(["open", "manual_review"]))
        rows = query.order_by(RequirementIssue.blocking.desc(), RequirementIssue.id.asc()).all()
        return [row.to_dict() for row in rows]

    def list_unresolved_blocking_issues(self, doc_id: int) -> List[Dict[str, Any]]:
        rows = (
            self.db.query(RequirementIssue)
            .filter(
                RequirementIssue.document_id == doc_id,
                RequirementIssue.blocking == 1,
                RequirementIssue.status.in_(["open", "manual_review"]),
            )
            .order_by(RequirementIssue.id.asc())
            .all()
        )
        return [row.to_dict() for row in rows]

    def list_requirement_revisions(self, doc_id: int) -> List[Dict[str, Any]]:
        document = self.get_document(doc_id)
        if not document:
            raise ValueError("需求文档不存在")
        rows = (
            self.db.query(RequirementRevision)
            .filter(RequirementRevision.document_id == doc_id)
            .order_by(RequirementRevision.revision_no.desc())
            .all()
        )
        return [row.to_dict() for row in rows]

    def get_requirement_issue(self, issue_id: int) -> Optional[RequirementIssue]:
        return self.db.query(RequirementIssue).filter(RequirementIssue.id == issue_id).first()

    def modify_issue(
        self,
        issue_id: int,
        revised_excerpt: str,
        *,
        operator: str = "tester",
        reason: str = "",
    ) -> Dict[str, Any]:
        issue = self.get_requirement_issue(issue_id)
        if not issue:
            raise ValueError("问题项不存在")
        return self._revise_issue(
            issue,
            revised_excerpt,
            operator=operator,
            reason=reason,
            action_type="modify",
        )

    def accept_ai_suggestion(
        self,
        issue_id: int,
        *,
        operator: str = "tester",
        reason: str = "",
    ) -> Dict[str, Any]:
        issue = self.get_requirement_issue(issue_id)
        if not issue:
            raise ValueError("问题项不存在")
        suggestion = (issue.suggestion or "").strip()
        if not suggestion:
            raise ValueError("问题项没有可采纳的 AI 修改建议")
        return self._revise_issue(
            issue,
            suggestion,
            operator=operator,
            reason=reason or "采纳 AI 修改建议",
            action_type="accept_ai_suggestion",
        )

    def _revise_issue(
        self,
        issue: RequirementIssue,
        revised_excerpt: str,
        *,
        operator: str,
        reason: str,
        action_type: str,
    ) -> Dict[str, Any]:
        document = self.get_document(issue.document_id)
        if not document or document.status != "active":
            raise ValueError("需求文档不存在")
        if not revised_excerpt.strip():
            raise ValueError("修订内容不能为空")

        original_excerpt = issue.source_excerpt or ""
        current_content = document.extracted_content or ""
        if original_excerpt and original_excerpt in current_content:
            next_content = current_content.replace(original_excerpt, revised_excerpt, 1)
        else:
            next_content = current_content.rstrip() + "\n\n" + revised_excerpt

        revision_no = self._next_revision_no(document.id)
        revision = RequirementRevision(
            document_id=document.id,
            issue_id=issue.id,
            revision_no=revision_no,
            original_excerpt=original_excerpt,
            revised_excerpt=revised_excerpt,
            full_content=next_content,
            diff_summary=json.dumps(
                {
                    "issue_id": issue.id,
                    "action": "modify",
                    "reason": reason,
                    "original_excerpt": original_excerpt,
                    "revised_excerpt": revised_excerpt,
                },
                ensure_ascii=False,
            ),
            created_by=operator,
        )
        self.db.add(revision)
        document.extracted_content = next_content
        document.parse_status = "unparsed"
        issue.status = "modified"
        issue.issue_type = "已修改"
        issue.operator = operator
        issue.blocking = 0
        self._record_issue_action(issue, action_type, operator, reason, {"revision_no": revision_no})
        self._resolve_issue_todo(issue, "问题项已修改，等待重新检查")
        self.db.commit()
        self.db.refresh(revision)
        self.db.refresh(issue)
        recheck_result = self.recheck_document(document.id, operator=operator)
        self.db.refresh(revision)
        self.db.refresh(issue)
        return {"issue": issue.to_dict(), "revision": revision.to_dict(), "document": recheck_result["document"]}

    def ignore_issue(
        self,
        issue_id: int,
        reason: str,
        *,
        operator: str = "tester",
        risk_accepted: bool = False,
    ) -> Dict[str, Any]:
        issue = self.get_requirement_issue(issue_id)
        if not issue:
            raise ValueError("问题项不存在")
        if not reason or len(reason.strip()) < 10:
            raise ValueError("忽略问题项必须填写不少于 10 字的原因")
        if issue.blocking and not risk_accepted:
            raise ValueError("阻断级问题不能无风险留痕忽略，请确认风险后再提交")

        issue.status = "ignored"
        issue.issue_type = "误报 / 已忽略"
        issue.ignored_reason = reason
        issue.operator = operator
        if risk_accepted:
            issue.blocking = 0
        
        # 同步更新关联的需求点状态
        if issue.requirement_item_id:
            item = self.db.query(RequirementItem).filter(RequirementItem.id == issue.requirement_item_id).first()
            if item:
                item.need_review = 0
                item.confirmed = 1

        self._record_issue_action(issue, "ignore", operator, reason, {"risk_accepted": risk_accepted})
        self._resolve_issue_todo(issue, "问题项已忽略并完成风险留痕")
        self._refresh_document_review_gate(issue.document_id)
        self.db.commit()
        self.db.refresh(issue)
        return issue.to_dict()

    def mark_issue_manual_review(self, issue_id: int, reason: str = "", operator: str = "tester") -> Dict[str, Any]:
        issue = self.get_requirement_issue(issue_id)
        if not issue:
            raise ValueError("问题项不存在")
        issue.status = "manual_review"
        issue.issue_type = "待确认"
        issue.operator = operator
        self._record_issue_action(issue, "manual_review", operator, reason, {})
        self.db.commit()
        self.db.refresh(issue)
        return issue.to_dict()

    def resolve_issue(self, issue_id: int, reason: str = "", operator: str = "tester") -> Dict[str, Any]:
        issue = self.get_requirement_issue(issue_id)
        if not issue:
            raise ValueError("问题项不存在")
        if issue.blocking and issue.status in {"open", "manual_review"}:
            raise ValueError("阻断级问题不能直接标记解决，请先修改后重新检查，或忽略并完成风险留痕")
        issue.status = "resolved"
        issue.issue_type = "已解决"
        issue.blocking = 0
        issue.operator = operator
        issue.resolved_at = sql_func.now()

        # 同步更新关联的需求点状态
        if issue.requirement_item_id:
            item = self.db.query(RequirementItem).filter(RequirementItem.id == issue.requirement_item_id).first()
            if item:
                item.need_review = 0
                item.confirmed = 1

        self._record_issue_action(issue, "resolve", operator, reason, {})
        self._resolve_issue_todo(issue, "问题项已解决")
        self._refresh_document_review_gate(issue.document_id)
        self.db.commit()
        self.db.refresh(issue)
        return issue.to_dict()

    def _refresh_document_review_gate(self, doc_id: int) -> None:
        document = self.get_document(doc_id)
        if not document or document.parse_status != "check_failed":
            return
        if not self.has_unresolved_blocking_issues(doc_id):
            document.parse_status = "pending_review"

    def recheck_document(self, doc_id: int, operator: str = "tester") -> Dict[str, Any]:
        document = self.get_document(doc_id)
        if not document or document.status != "active":
            raise ValueError("需求文档不存在")
        result = self.trigger_parse(doc_id)
        for issue in self.db.query(RequirementIssue).filter(RequirementIssue.document_id == doc_id).all():
            self._record_issue_action(issue, "recheck", operator, "重新检查文档", {})
        self.db.commit()
        return result

    def recheck_issue(self, issue_id: int, operator: str = "tester") -> Dict[str, Any]:
        issue = self.get_requirement_issue(issue_id)
        if not issue:
            raise ValueError("问题项不存在")
        self._record_issue_action(issue, "recheck", operator, "重新检查问题项关联文档", {"issue_id": issue.id})
        self.db.commit()
        return self.recheck_document(issue.document_id, operator=operator)

    def has_unresolved_blocking_issues(self, doc_id: int) -> bool:
        return bool(self.list_unresolved_blocking_issues(doc_id))

    async def create_document_version(self, doc_id: int, file: UploadFile, created_by: str = "") -> RequirementDocument:
        """上传新版本文档，保留与上一版本的父子关系"""
        parent = self.get_document(doc_id)
        if not parent or parent.status != "active":
            raise ValueError("需求文档不存在")
        next_version = (parent.version_no or 1) + 1
        new_doc = await self.save_document(
            file=file,
            title=f"{parent.title} v{next_version}",
            category=parent.category,
            module=parent.module,
            dependency_scope=parent.dependency_scope,
            dependency_notes=parent.dependency_notes,
            tags=parent.tags,
            created_by=created_by or parent.created_by,
            tree_node_id=parent.tree_node_id,
            project_id=parent.project_id,
        )
        new_doc.parent_document_id = parent.id
        new_doc.version_no = next_version
        self.db.commit()
        self.db.refresh(new_doc)
        return new_doc

    def get_tree_path(self, doc_id: int) -> Dict:
        document = self.get_document(doc_id)
        if not document:
            raise ValueError("需求文档不存在")
        if not document.tree_node_id:
            return {"document_id": doc_id, "path": [], "path_label": "未挂载"}
        nodes = {
            node.id: node
            for node in self.db.query(RequirementTreeNode).filter(RequirementTreeNode.is_active == 1).all()
        }
        path = []
        current_id = document.tree_node_id
        while current_id and current_id in nodes:
            node = nodes[current_id]
            path.append({"id": node.id, "name": node.name, "node_type": node.node_type})
            current_id = node.parent_id
        path.reverse()
        label = " / ".join(f"{item['node_type']}:{item['name']}" for item in path)
        return {"document_id": doc_id, "path": path, "path_label": label or "未挂载"}

    def list_requirement_items(self, doc_id: int) -> List[Dict]:
        document = self.get_document(doc_id)
        if not document:
            raise ValueError("需求文档不存在")
        return [item.to_dict() for item in document.requirement_items.all()]

    def _build_history_diff(self, document: RequirementDocument, new_points: List[Dict[str, Any]]) -> Dict[str, Any]:
        if not document.module:
            return {"diff_items": [], "summary": {}, "conflict_count": 0}
        old_docs = (
            self.db.query(RequirementDocument)
            .filter(
                RequirementDocument.id != document.id,
                RequirementDocument.module == document.module,
                RequirementDocument.status == "active",
                RequirementDocument.parse_status == "stored",
            )
            .order_by(RequirementDocument.id.desc())
            .limit(3)
            .all()
        )
        old_points: List[Dict[str, Any]] = []
        for old_doc in old_docs:
            for item in old_doc.requirement_items:
                old_points.append(item.to_dict())
        if not old_points:
            return {"diff_items": [], "summary": {}, "conflict_count": 0}
        skill = RequirementDiffSkill()
        result = skill.run({"old_requirements": old_points, "new_requirements": new_points})
        data = result.get("data", {})
        diff_items = data.get("diff_items", [])
        conflict_count = sum(1 for item in diff_items if item.get("change_type") == "modified")
        return {
            "diff_items": diff_items,
            "summary": data.get("summary", {}),
            "conflict_count": conflict_count,
        }

    def _build_coupling_results(self, document: RequirementDocument, points: List[Dict[str, Any]]) -> Dict[str, Any]:
        items = []
        for point in points:
            scope = point.get("dependency_scope") or []
            if isinstance(scope, str):
                scope = [item.strip() for item in scope.split(",") if item.strip()]
            notes = point.get("dependency_notes") or point.get("dependency_consideration") or ""
            if scope or notes:
                items.append(
                    {
                        "requirement_no": point.get("requirement_no") or "",
                        "title": point.get("title") or point.get("content", "")[:40],
                        "dependency_scope": scope,
                        "dependency_notes": notes,
                        "source_excerpt": point.get("source_text") or "",
                    }
                )
        if document.dependency_scope or document.dependency_notes:
            items.append(
                {
                    "requirement_no": "DOCUMENT",
                    "title": document.title,
                    "dependency_scope": [item.strip() for item in (document.dependency_scope or "").split(",") if item.strip()],
                    "dependency_notes": document.dependency_notes or "",
                    "source_excerpt": document.extracted_content[:200] if document.extracted_content else "",
                }
            )
        return {
            "has_coupling": bool(items),
            "items": items,
            "summary": "发现业务耦合依赖，请在入库前确认上下游影响。" if items else "未发现显式业务耦合依赖。",
        }

    def _run_bound_parse_skill(self, document: RequirementDocument, document_type: str, parser_skill: str = "") -> Dict[str, Any]:
        content = document.extracted_content or ""
        if document_type == "api_document":
            result = ApiDocumentParseSkill().run({"content": content, "title": document.title})
            data = result.get("data", {})
            if not result.get("success"):
                raise ValueError(result.get("message") or "接口文档解析失败")
            return {
                "requirement_points": data.get("requirement_points", []),
                "issues": data.get("issues", []),
                "skill_output": data.get("skill_output", data),
            }
        if document_type == "api_log":
            result = ApiLogParseSkill().run({"content": content})
            data = result.get("data", {})
            return {
                "requirement_points": [],
                "issues": data.get("issues", []),
                "entries": data.get("entries", []),
                "skill_output": data,
            }

        result = RequirementParseSkill().run(
            {
                "content": content,
                "module": document.module or "",
                "version": "",
            }
        )
        data = result.get("data", {})
        if not result.get("success"):
            raise ValueError(result.get("message") or "AI 需求解析失败")
        issues = data.get("issues", [])
        return {
            "requirement_points": data.get("requirement_points", []),
            "issues": issues,
            "skill_output": data,
        }

    def _build_api_requirement_points(self, analysis: Dict[str, Any]) -> List[Dict[str, Any]]:
        points = []
        for index, endpoint in enumerate(analysis.get("endpoints", []), start=1):
            methods = endpoint.get("methods") or []
            method_text = "/".join(methods) if methods else "METHOD"
            path = endpoint.get("path", "")
            description = endpoint.get("description") or "接口定义"
            source_text = f"{method_text} {path} {description}".strip()
            points.append(
                {
                    "requirement_no": f"API-{index:03d}",
                    "module": endpoint.get("category") or "接口文档",
                    "title": f"{method_text} {path} - {description}"[:120],
                    "content": f"接口 {method_text} {path} 应符合文档定义：{description}",
                    "type": "api_contract",
                    "priority": "P1",
                    "risk_level": "medium",
                    "dependency_scope": ["api"],
                    "version": "",
                    "source_text": source_text,
                    "need_review": not endpoint.get("assertions"),
                }
            )
        return points

    def _build_parse_issues(self, points: List[Dict[str, Any]], document: RequirementDocument) -> List[Dict[str, str]]:
        issues: List[Dict[str, str]] = []
        if not (document.extracted_content or "").strip():
            issues.append({"type": "missing", "severity": "high", "message": "文档未提取到有效文本内容"})
        if not points:
            issues.append({"type": "missing", "severity": "high", "message": "未解析出结构化需求点"})
        for point in points:
            title = point.get("title") or point.get("content", "")[:15]
            # 去除可能包含的 Markdown 标题符号，让显示更简洁
            title = str(title).lstrip("# ").strip()
            if point.get("need_review"):
                issues.append(
                    {
                        "type": "ambiguous",
                        "message": f"需求点《{title}》 ({point.get('requirement_no')}) 存在待确认表述",
                        "requirement_no": point.get("requirement_no"),
                        "requirement_title": title,
                    }
                )
            if len(point.get("content", "")) < 6:
                issues.append(
                    {
                        "type": "missing",
                        "message": f"需求点《{title}》 ({point.get('requirement_no')}) 描述过短，信息不足",
                        "requirement_no": point.get("requirement_no"),
                        "requirement_title": title,
                    }
                )
        return issues

    def _build_need_review_issues(self, points: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        issues: List[Dict[str, Any]] = []
        for point in points:
            if not point.get("need_review"):
                continue
            requirement_no = point.get("requirement_no") or ""
            title = str(point.get("title") or point.get("content", "")[:40]).lstrip("# ").strip()
            source_text = point.get("source_text") or point.get("content") or title
            issues.append(
                {
                    "type": "ambiguous",
                    "severity": "medium",
                    "requirement_no": requirement_no,
                    "requirement_title": title,
                    "source_excerpt": source_text,
                    "message": f"需求点《{title}》({requirement_no}) 需要人工复核后才能确认入库",
                    "suggestion": "请补充接口断言、响应字段或验收口径后重新检查；如确认无误，可在问题项中完成处理留痕。",
                    "ai_reason": point.get("review_reason") or "解析结果缺少足够的可验证断言或验收口径",
                }
            )
        return issues

    def _sync_requirement_issues(self, document: RequirementDocument, issues: List[Dict[str, Any]]) -> None:
        """将解析结果中的 issues 标准化落库，形成可处理的问题项。"""
        manual_review_issues = (
            self.db.query(RequirementIssue)
            .filter(
                RequirementIssue.document_id == document.id,
                RequirementIssue.status == "manual_review",
            )
            .all()
        )
        manual_review_by_key = {
            (issue.source_location or "", issue.source_excerpt or ""): issue
            for issue in manual_review_issues
        }
        stale_issue_ids = [
            row[0]
            for row in self.db.query(RequirementIssue.id)
            .filter(
                RequirementIssue.document_id == document.id,
                RequirementIssue.status == "open",
            )
            .all()
        ]
        for issue_id in stale_issue_ids:
            self._resolve_issue_todo_by_id(issue_id, "需求重新解析后旧问题项已刷新，自动核销旧待办")
        self.db.query(RequirementIssue).filter(
            RequirementIssue.document_id == document.id,
            RequirementIssue.status == "open",
        ).delete(synchronize_session=False)
        requirement_items = {
            item.requirement_no: item
            for item in self.db.query(RequirementItem)
            .filter(RequirementItem.document_id == document.id)
            .all()
        }
        for raw_issue in issues:
            normalized = self._normalize_requirement_issue(raw_issue, document, requirement_items)
            key = (normalized.get("source_location") or "", normalized.get("source_excerpt") or "")
            manual_issue = manual_review_by_key.pop(key, None)
            if manual_issue:
                manual_issue.requirement_item_id = normalized.get("requirement_item_id")
                manual_issue.severity = normalized.get("severity", manual_issue.severity)
                manual_issue.blocking = normalized.get("blocking", manual_issue.blocking)
                manual_issue.source_excerpt = normalized.get("source_excerpt", manual_issue.source_excerpt)
                manual_issue.title = normalized.get("title", manual_issue.title)
                manual_issue.message = normalized.get("message", manual_issue.message)
                manual_issue.suggestion = normalized.get("suggestion", manual_issue.suggestion)
                manual_issue.ai_reason = normalized.get("ai_reason", manual_issue.ai_reason)
                manual_issue.impact_scope = normalized.get("impact_scope", manual_issue.impact_scope)
                continue
            self.db.add(RequirementIssue(**normalized))
        self.db.flush()

        for stale_manual_issue in manual_review_by_key.values():
            stale_manual_issue.status = "resolved"
            stale_manual_issue.issue_type = "已解决"
            stale_manual_issue.blocking = 0
            stale_manual_issue.resolved_at = sql_func.now()
            self._record_issue_action(stale_manual_issue, "recheck_resolved", "system", "重新检查后问题已消失", {})
            self._resolve_issue_todo(stale_manual_issue, "重新检查后问题已消失")

        # 更新已保留的历史问题项（已解决、已忽略、已修改）关联的新 RequirementItem ID
        preserved_issues = self.db.query(RequirementIssue).filter(
            RequirementIssue.document_id == document.id,
            RequirementIssue.status.in_(["resolved", "ignored", "modified"])
        ).all()
        for issue in preserved_issues:
            if issue.source_location and issue.source_location in requirement_items:
                issue.requirement_item_id = requirement_items[issue.source_location].id
        self.db.flush()

    def _normalize_requirement_issue(
        self,
        issue: Dict[str, Any],
        document: RequirementDocument,
        requirement_items: Dict[str, RequirementItem],
    ) -> Dict[str, Any]:
        raw_type = str(issue.get("type") or "ambiguous")
        raw_severity = str(issue.get("severity") or "").lower()
        requirement_no = issue.get("requirement_no") or ""
        item = requirement_items.get(requirement_no)
        severity = self._map_issue_severity(raw_type, raw_severity)
        issue_type = self._map_issue_type(raw_type, severity)
        blocking = 1 if severity in {"阻断", "高"} and issue_type != "待优化" else 0
        source_excerpt = (
            issue.get("source_excerpt")
            or issue.get("source_text")
            or (item.source_text if item else "")
            or ""
        )
        return {
            "document_id": document.id,
            "requirement_item_id": item.id if item else None,
            "issue_type": issue_type,
            "severity": severity,
            "status": "open",
            "blocking": blocking,
            "source_location": requirement_no or issue.get("source_location", ""),
            "source_excerpt": source_excerpt,
            "title": issue.get("requirement_title") or issue.get("title") or issue_type,
            "message": issue.get("message") or issue.get("reason") or "需求检查发现待处理问题",
            "suggestion": issue.get("suggestion") or self._default_issue_suggestion(issue_type),
            "ai_reason": issue.get("ai_reason") or issue.get("message") or "",
            "impact_scope": issue.get("impact_scope") or document.module or "",
        }

    def _map_issue_severity(self, raw_type: str, raw_severity: str) -> str:
        if raw_severity in {"blocker", "blocking", "critical", "阻断"}:
            return "阻断"
        if raw_severity in {"high", "高"} or raw_type in {"missing", "conflict", "runtime_error"}:
            return "高"
        if raw_severity in {"low", "低"}:
            return "低"
        return "中"

    def _map_issue_type(self, raw_type: str, severity: str) -> str:
        if raw_type in {"logic_gap", "logic_error", "logical_gap", "logical_error", "逻辑漏洞"}:
            return "逻辑漏洞"
        if raw_type in {"format_error", "format", "schema_error", "格式问题"}:
            return "格式问题"
        if raw_type in {"missing", "conflict", "runtime_error"}:
            return "待修改"
        if raw_type in {"ambiguous", "need_review"}:
            return "待确认"
        if severity in {"低", "中"}:
            return "待优化"
        return "待确认"

    def _default_issue_suggestion(self, issue_type: str) -> str:
        if issue_type == "逻辑漏洞":
            return "请补充缺失的业务分支、异常路径或状态流转后重新检查。"
        if issue_type == "格式问题":
            return "请按需求模板补齐标题、验收标准、表格或字段格式后重新检查。"
        if issue_type == "待修改":
            return "请在平台内修改对应原文片段，或补充完整后重新检查。"
        if issue_type == "待优化":
            return "可优化描述后重新检查；如确认不影响入库，可填写原因忽略。"
        return "请产品或研发确认业务含义后再入库。"

    def _record_issue_action(
        self,
        issue: RequirementIssue,
        action_type: str,
        operator: str,
        reason: str,
        payload: Dict[str, Any],
    ) -> None:
        self.db.add(
            RequirementIssueAction(
                issue_id=issue.id,
                document_id=issue.document_id,
                action_type=action_type,
                operator=operator or "system",
                reason=reason or "",
                payload=json.dumps(payload or {}, ensure_ascii=False),
            )
        )

    def _sync_issue_todos(self, document: RequirementDocument) -> None:
        from app.services.todo_service import TodoService

        todo_service = TodoService(self.db)
        blocking_issues = (
            self.db.query(RequirementIssue)
            .filter(
                RequirementIssue.document_id == document.id,
                RequirementIssue.blocking == 1,
                RequirementIssue.status.in_(["open", "manual_review"]),
            )
            .all()
        )
        for issue in blocking_issues:
            todo_service.register_todo(
                source_type="REQ_ISSUE_BLOCKING",
                source_id=issue.id,
                title=f"需求阻断问题待处理: 《{document.title}》",
                description=f"{issue.issue_type} / {issue.severity}: {issue.message}",
                importance="high",
                risk_level="high",
            )

    def _resolve_issue_todo(self, issue: RequirementIssue, reason: str) -> None:
        self._resolve_issue_todo_by_id(issue.id, reason)

    def _resolve_issue_todo_by_id(self, issue_id: int, reason: str) -> None:
        from app.services.todo_service import TodoService

        TodoService(self.db).resolve_todo("REQ_ISSUE_BLOCKING", issue_id, reason=reason)

    def _next_revision_no(self, doc_id: int) -> int:
        latest = (
            self.db.query(RequirementRevision)
            .filter(RequirementRevision.document_id == doc_id)
            .order_by(RequirementRevision.revision_no.desc())
            .first()
        )
        return (latest.revision_no + 1) if latest else 1

    def list_categories(self) -> List[str]:
        rows = self.db.query(RequirementDocument.category).distinct().all()
        return [row[0] for row in rows if row[0]]

    def get_document(self, doc_id: int) -> Optional[RequirementDocument]:
        return self.db.query(RequirementDocument).filter(RequirementDocument.id == doc_id).first()

    def build_generation_context(self, doc_id: int) -> Dict:
        """构造 AI 生成 case 所需的上下文，包含同分类文档辅助信息"""
        self.assert_can_generate_cases(doc_id)
        document = self.get_document(doc_id)
        if not document:
            raise ValueError("需求文档不存在")

        related_docs = (
            self.db.query(RequirementDocument)
            .filter(
                RequirementDocument.id != doc_id,
                RequirementDocument.category == document.category,
                RequirementDocument.status == "active",
            )
            .limit(5)
            .all()
        )

        related_context = [
            {
                "title": item.title,
                "module": item.module,
                "dependency_scope": item.dependency_scope,
                "summary": item.ai_summary,
                "content_excerpt": item.extracted_content[:800],
            }
            for item in related_docs
        ]
        requirement_points = [
            {
                "id": item.id,
                "requirement_no": item.requirement_no,
                "module": document.module or document.category or "",
                "title": item.title,
                "content": item.content,
                "type": item.item_type,
                "priority": item.priority,
                "source_text": item.source_text,
            }
            for item in (
                self.db.query(RequirementItem)
                .filter(
                    RequirementItem.document_id == doc_id,
                    RequirementItem.confirmed == 1,
                    RequirementItem.need_review == 0,
                )
                .order_by(RequirementItem.id.asc())
                .all()
            )
        ]

        return {
            "document": document.to_dict(),
            "requirement_points": requirement_points,
            "related_documents": related_context,
            "analysis": self.analyze_document(doc_id),
        }

    def update_ai_summary(self, doc_id: int, summary: str) -> None:
        document = self.get_document(doc_id)
        if not document:
            return
        if self._is_unusable_ai_summary(summary):
            return
        document.ai_summary = summary
        self.db.commit()

    def _is_unusable_ai_summary(self, summary: Optional[str]) -> bool:
        if not summary:
            return False
        stripped = summary.strip()
        return stripped.startswith(("[AI 服务未配置]", "[AI 调用失败]"))

    def _build_rule_parse_summary(self, points: List[Dict[str, Any]], issues: List[Dict[str, str]]) -> str:
        if not points:
            return "规则解析未提取到结构化需求点，请检查原文内容是否完整。"

        modules = Counter(point.get("module") or "未分类模块" for point in points)
        priorities = Counter(point.get("priority") or "P1" for point in points)
        module_text = "、".join(f"{name} {count} 条" for name, count in modules.most_common(3))
        priority_text = "、".join(f"{name} {count} 条" for name, count in priorities.most_common())
        issue_text = f"发现 {len(issues)} 个待处理问题" if issues else "未发现阻断入库的问题"
        return f"规则解析完成：共提取 {len(points)} 条需求点；主要模块：{module_text}；优先级分布：{priority_text}；{issue_text}。"

    def analyze_document(self, doc_id: int) -> Dict[str, Any]:
        """读取接口文档结构化分析结果。"""
        document = self.get_document(doc_id)
        if not document:
            raise ValueError("需求文档不存在")
        analysis = None
        if document.parse_result:
            try:
                payload = json.loads(document.parse_result) if isinstance(document.parse_result, str) else document.parse_result
                analysis = payload.get("skill_output") or {}
            except (TypeError, json.JSONDecodeError):
                analysis = None
        if not analysis:
            result = ApiDocumentParseSkill().run({"content": document.extracted_content or "", "title": document.title})
            if not result.get("success"):
                raise ValueError(result.get("message") or "接口文档分析失败")
            analysis = result.get("data", {}).get("skill_output", {})
        analysis["document_id"] = document.id
        analysis["title"] = document.title
        return analysis

    def generate_api_case_drafts(self, doc_id: int, options: Dict[str, Any]) -> Dict[str, Any]:
        """基于结构化解析结果生成可执行 API case 草稿。"""
        self.assert_can_generate_cases(doc_id)
        document = self.get_document(doc_id)
        if not document:
            raise ValueError("需求文档不存在")

        analysis = self.analyze_document(doc_id)
        mode = options.get("generation_mode") or "status_codes"
        case_count = int(options.get("case_count") or 5)
        cases = self._build_api_cases(document.to_dict(), analysis, mode, case_count)

        return {
            "summary": self._build_generation_summary(analysis, mode, len(cases)),
            "test_focus": analysis.get("test_focus", []),
            "analysis": {
                "base_url": analysis.get("base_url", ""),
                "endpoint_count": analysis.get("stats", {}).get("endpoint_count", 0),
                "status_case_count": len(analysis.get("status_scenarios", [])),
                "warnings": analysis.get("warnings", []),
            },
            "cases": cases,
        }

    def _build_api_cases(
        self,
        document: Dict[str, Any],
        analysis: Dict[str, Any],
        mode: str,
        case_count: int,
    ) -> List[Dict[str, Any]]:
        if mode == "status_codes":
            scenarios = analysis.get("status_scenarios", [])
            selected = scenarios[: max(case_count, 4)]
            cases = [self._status_scenario_to_case(document, scenario) for scenario in selected]
            return self._normalize_generated_cases(document, cases, "api")

        endpoints = analysis.get("endpoints", [])
        if mode == "smoke":
            endpoints = endpoints[:case_count]
        else:
            endpoints = endpoints[: max(case_count, len(endpoints))]

        cases = [self._endpoint_to_case(document, analysis.get("base_url", ""), endpoint) for endpoint in endpoints]
        if not cases and analysis.get("status_scenarios"):
            cases = [self._status_scenario_to_case(document, item) for item in analysis["status_scenarios"][:case_count]]
        cases = cases[:case_count] if mode != "full" else cases
        return self._normalize_generated_cases(document, cases, "api")

    def _normalize_generated_cases(
        self,
        document: Dict[str, Any],
        cases: List[Dict[str, Any]],
        case_type: str,
    ) -> List[Dict[str, Any]]:
        result = CaseGenerateSkill().run(
            {
                "document": document,
                "cases": cases,
                "case_type": case_type,
            }
        )
        if result.get("success"):
            return result.get("data", {}).get("cases", cases)
        return cases

    def _status_scenario_to_case(self, document: Dict[str, Any], scenario: Dict[str, Any]) -> Dict[str, Any]:
        expected_status = scenario.get("expected_status")
        expected_statuses = scenario.get("expected_statuses") or []
        path = scenario["path"]
        status_label = expected_status if expected_status else "/".join(str(code) for code in expected_statuses)
        return {
            "name": f"httpbin 状态码 {status_label} 校验",
            "category": "positive" if expected_status and expected_status < 400 else "negative",
            "priority": "P0" if expected_status in {200, 404, 500} or len(expected_statuses) > 1 else "P1",
            "case_type": "api",
            "method": scenario.get("method", "GET"),
            "path": path,
            "base_url": scenario.get("base_url", ""),
            "request_data": {"method": scenario.get("method", "GET"), "path": path, "params": {}, "body": None},
            "request_headers": {},
            "expected_status": expected_status,
            "expected_statuses": expected_statuses,
            "expected_body": {},
            "expected_contains": [],
            "assertions": [{"type": "status_code", "expected": expected_status}] if expected_status else [{"type": "status_in", "expected": expected_statuses}],
            "precondition": f"服务 {scenario.get('base_url') or 'httpbin'} 可访问",
            "steps": [
                f"发送 {scenario.get('method', 'GET')} 请求到 {path}",
                f"校验响应状态码为 {status_label}",
            ],
            "expected_result": f"接口返回状态码 {status_label}",
            "dependency_consideration": f"来源文档：{document.get('title')}；状态码接口不依赖业务数据。",
            "source_excerpt": scenario.get("description", ""),
            "selected": True,
        }

    def _endpoint_to_case(self, document: Dict[str, Any], base_url: str, endpoint: Dict[str, Any]) -> Dict[str, Any]:
        method = (endpoint.get("methods") or ["GET"])[0]
        path = endpoint.get("path", "")
        assertions = endpoint.get("assertions") or [{"type": "status_code", "expected": 200}]
        expected_status = self._expected_status_from_assertions(assertions)
        return {
            "name": f"{method} {path} 冒烟校验",
            "category": "positive",
            "priority": "P1",
            "case_type": "api",
            "method": method,
            "path": path,
            "base_url": base_url,
            "request_data": {"method": method, "path": path, "params": {}, "body": None},
            "request_headers": {},
            "expected_status": expected_status,
            "expected_statuses": [expected_status] if expected_status else [],
            "expected_body": {},
            "expected_contains": [],
            "assertions": assertions,
            "precondition": f"服务 {base_url or '被测服务'} 可访问",
            "steps": [f"发送 {method} 请求到 {path}", "校验响应状态和关键字段"],
            "expected_result": endpoint.get("description") or "接口响应符合文档描述",
            "dependency_consideration": f"来源文档：{document.get('title')}",
            "source_excerpt": endpoint.get("description", ""),
            "selected": True,
        }

    def _expected_status_from_assertions(self, assertions: List[Dict[str, Any]]) -> Optional[int]:
        for assertion in assertions:
            if assertion.get("type") == "status_code" and isinstance(assertion.get("expected"), int):
                return assertion["expected"]
        return 200

    def _build_generation_summary(self, analysis: Dict[str, Any], mode: str, count: int) -> str:
        mode_label = {"status_codes": "状态码专项", "smoke": "接口主流程", "full": "全量接口冒烟"}.get(mode, mode)
        return f"已基于结构化解析生成 {count} 条 API case 草稿，模式：{mode_label}。"

    def _extract_text(self, filepath: Path, suffix: str) -> str:
        if suffix in {".md", ".txt"}:
            return filepath.read_text(encoding="utf-8", errors="ignore")

        if suffix == ".docx":
            try:
                from docx import Document
            except ImportError:
                return "[缺少 python-docx，暂无法解析 docx 内容]"
            doc = Document(str(filepath))
            return "\n".join(p.text for p in doc.paragraphs if p.text).strip()

        if suffix == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError:
                return "[缺少 pypdf，暂无法解析 pdf 内容]"

            reader = PdfReader(str(filepath))
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts).strip()

        return ""

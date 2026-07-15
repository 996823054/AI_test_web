"""Requirement domain service module (split from RequirementDocService)."""

import hashlib
import json
import os
import uuid
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List, Optional
from fastapi import UploadFile
from sqlalchemy.sql import func as sql_func
from sqlalchemy.orm import Session
from app.config import settings
from app.models.requirement_document import RequirementDocument
from app.models.requirement_issue import RequirementIssue, RequirementIssueAction, RequirementRevision
from app.models.requirement_item import RequirementItem
from app.models.requirement_tree_node import RequirementTreeNode
from app.models.test_case import TestCase
from app.models.changelog import Changelog
from app.models.api_info import APIInfo
from app.services.requirement_tree_service import RequirementTreeService
from app.skills.api_document_parse_skill import ApiDocumentParseSkill
from app.skills.api_log_parse_skill import ApiLogParseSkill
from app.skills.case_generate_skill import CaseGenerateSkill
from app.skills.document_parse_router_skill import DocumentParseRouterSkill
from app.skills.requirement_diff_skill import RequirementDiffSkill
from app.skills.requirement_parse_skill import RequirementParseSkill

class RequirementDocumentModule:
    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}

    PARSE_STATUSES = {"unparsed", "parsing", "pending_review", "check_failed", "stored"}

    def __init__(self, db: Session):
            self.db = db
            self.upload_dir = Path(settings.DOCUMENT_UPLOAD_DIR)
            self.upload_dir.mkdir(parents=True, exist_ok=True)

    async def save_document(
            self,
            file: UploadFile,
            title: str,
            category: str,
            module: str,
            dependency_scope: str,
            dependency_notes: str,
            tags: str,
            created_by: str,
            tree_node_id: Optional[int] = None,
            project_id: Optional[int] = None,
        ) -> RequirementDocument:
            """保存上传文档并提取文本"""
            if tree_node_id is not None:
                node = RequirementTreeService(self.db).get_node(tree_node_id)
                if not node:
                    raise ValueError("树节点不存在")
            suffix = Path(file.filename or "").suffix.lower()
            if suffix not in self.ALLOWED_EXTENSIONS:
                raise ValueError("仅支持 pdf、docx、md、txt 文件")

            file_id = uuid.uuid4().hex
            saved_name = f"{file_id}{suffix}"
            saved_path = self.upload_dir / saved_name

            content = await file.read()
            file_hash = hashlib.sha256(content).hexdigest()
            duplicate = (
                self.db.query(RequirementDocument)
                .filter(
                    RequirementDocument.file_hash == file_hash,
                    RequirementDocument.status == "active",
                )
                .first()
            )
            if duplicate:
                raise ValueError(f"检测到重复文件，已存在文档《{duplicate.title}》")

            doc_title = title or Path(file.filename or saved_name).stem
            title_duplicate = (
                self.db.query(RequirementDocument)
                .filter(
                    RequirementDocument.title == doc_title,
                    RequirementDocument.status == "active",
                )
                .first()
            )
            if title_duplicate:
                raise ValueError(f"已存在同名文档《{title_duplicate.title}》，请修改标题或归档旧文档")

            with open(saved_path, "wb") as output:
                output.write(content)

            extracted_content = self._extract_text(saved_path, suffix)

            document = RequirementDocument(
                title=doc_title,
                file_name=file.filename or saved_name,
                file_type=suffix.replace(".", ""),
                file_path=str(saved_path),
                file_size=len(content),
                file_hash=file_hash,
                project_id=project_id,
                tree_node_id=tree_node_id,
                category=category or "未分类",
                module=module or "",
                dependency_scope=dependency_scope or "",
                dependency_notes=dependency_notes or "",
                tags=tags or "",
                extracted_content=extracted_content,
                parse_status="unparsed",
                created_by=created_by or "",
            )
            self.db.add(document)
            self.db.commit()
            self.db.refresh(document)
            return document

    def list_documents(
            self,
            category: Optional[str] = None,
            keyword: Optional[str] = None,
            tree_node_id: Optional[int] = None,
            status: str = "active",
            page: int = 1,
            page_size: int = 20,
        ) -> Dict:
            query = self.db.query(RequirementDocument).filter(RequirementDocument.status == status)
            if category:
                query = query.filter(RequirementDocument.category == category)
            if tree_node_id is not None:
                query = query.filter(RequirementDocument.tree_node_id == tree_node_id)
            if keyword:
                query = query.filter(
                    (RequirementDocument.title.contains(keyword))
                    | (RequirementDocument.file_name.contains(keyword))
                    | (RequirementDocument.module.contains(keyword))
                )

            total = query.count()
            items = (
                query.order_by(RequirementDocument.created_at.desc())
                .offset((page - 1) * page_size)
                .limit(page_size)
                .all()
            )
            return {
                "total": total,
                "page": page,
                "page_size": page_size,
                "items": [item.to_dict() for item in items],
            }

    def update_document_meta(
            self,
            doc_id: int,
            title: Optional[str] = None,
            category: Optional[str] = None,
            module: Optional[str] = None,
        ) -> RequirementDocument:
            document = self.get_document(doc_id)
            if not document or document.status != "active":
                raise ValueError("需求文档不存在")
            if title is not None:
                document.title = title
            if category is not None:
                document.category = category
            if module is not None:
                document.module = module
            self.db.commit()
            self.db.refresh(document)
            return document

    def archive_document(self, doc_id: int) -> RequirementDocument:
            document = self.get_document(doc_id)
            if not document or document.status != "active":
                raise ValueError("需求文档不存在")
            document.status = "archived"
            self.db.commit()
            self.db.refresh(document)
            return document

    def soft_delete_document(self, doc_id: int) -> RequirementDocument:
            document = self.get_document(doc_id)
            if not document or document.status != "active":
                raise ValueError("需求文档不存在")
            document.status = "deleted"
            self.db.commit()
            self.db.refresh(document)
            return document

    def restore_document(self, doc_id: int) -> RequirementDocument:
            document = self.db.query(RequirementDocument).filter(RequirementDocument.id == doc_id).first()
            if not document or document.status not in {"archived", "deleted"}:
                raise ValueError("文档不在归档或回收站中")
            document.status = "active"
            self.db.commit()
            self.db.refresh(document)
            return document

    async def create_document_version(self, doc_id: int, file: UploadFile, created_by: str = "") -> RequirementDocument:
            """上传新版本文档，保留与上一版本的父子关系"""
            parent = self.get_document(doc_id)
            if not parent or parent.status != "active":
                raise ValueError("需求文档不存在")
            next_version = (parent.version_no or 1) + 1
            new_doc = await self.save_document(
                file=file,
                title=f"{parent.title} v{next_version}",
                category=parent.category,
                module=parent.module,
                dependency_scope=parent.dependency_scope,
                dependency_notes=parent.dependency_notes,
                tags=parent.tags,
                created_by=created_by or parent.created_by,
                tree_node_id=parent.tree_node_id,
                project_id=parent.project_id,
            )
            new_doc.parent_document_id = parent.id
            new_doc.version_no = next_version
            self.db.commit()
            self.db.refresh(new_doc)
            return new_doc

    def get_document_impact(self, doc_id: int) -> Dict[str, int]:
            document = self.get_document(doc_id)
            if not document:
                raise ValueError("需求文档不存在")
            item_count = self.db.query(RequirementItem).filter(RequirementItem.document_id == doc_id).count()
            case_count = (
                self.db.query(TestCase)
                .filter(TestCase.source_document_id == doc_id, TestCase.is_active == 1)
                .count()
            )
            api_ids = [
                row[0]
                for row in self.db.query(TestCase.api_id)
                .filter(TestCase.source_document_id == doc_id, TestCase.api_id.isnot(None))
                .distinct()
                .all()
                if row[0]
            ]
            change_record_count = self.db.query(Changelog).filter(Changelog.api_id.in_(api_ids)).count() if api_ids else 0
            return {
                "document_id": doc_id,
                "requirement_item_count": item_count,
                "case_count": case_count,
                "knowledge_fragment_count": item_count,
                "change_record_count": change_record_count,
            }

    def get_document_relations(self, doc_id: int) -> Dict[str, Any]:
            document = self.get_document(doc_id)
            if not document:
                raise ValueError("需求文档不存在")
            cases = (
                self.db.query(TestCase)
                .filter(TestCase.source_document_id == doc_id)
                .order_by(TestCase.id.asc())
                .all()
            )
            api_ids = sorted({case.api_id for case in cases if case.api_id})
            apis = self.db.query(APIInfo).filter(APIInfo.id.in_(api_ids)).all() if api_ids else []
            changes = self.db.query(Changelog).filter(Changelog.api_id.in_(api_ids)).all() if api_ids else []
            modules = sorted({api.module for api in apis if api.module})
            return {
                "document": document.to_dict(),
                "cases": [case.to_dict() for case in cases],
                "apis": [api.to_dict() for api in apis],
                "modules": modules,
                "changes": [change.to_dict() for change in changes],
            }

    def get_tree_path(self, doc_id: int) -> Dict:
            document = self.get_document(doc_id)
            if not document:
                raise ValueError("需求文档不存在")
            if not document.tree_node_id:
                return {"document_id": doc_id, "path": [], "path_label": "未挂载"}
            nodes = {
                node.id: node
                for node in self.db.query(RequirementTreeNode).filter(RequirementTreeNode.is_active == 1).all()
            }
            path = []
            current_id = document.tree_node_id
            while current_id and current_id in nodes:
                node = nodes[current_id]
                path.append({"id": node.id, "name": node.name, "node_type": node.node_type})
                current_id = node.parent_id
            path.reverse()
            label = " / ".join(f"{item['node_type']}:{item['name']}" for item in path)
            return {"document_id": doc_id, "path": path, "path_label": label or "未挂载"}

    def list_categories(self) -> List[str]:
            rows = self.db.query(RequirementDocument.category).distinct().all()
            return [row[0] for row in rows if row[0]]

    def get_document(self, doc_id: int) -> Optional[RequirementDocument]:
            return self.db.query(RequirementDocument).filter(RequirementDocument.id == doc_id).first()

    def update_ai_summary(self, doc_id: int, summary: str) -> None:
            document = self.get_document(doc_id)
            if not document:
                return
            if self._is_unusable_ai_summary(summary):
                return
            document.ai_summary = summary
            self.db.commit()

    def _is_unusable_ai_summary(self, summary: Optional[str]) -> bool:
            if not summary:
                return False
            stripped = summary.strip()
            return stripped.startswith(("[AI 服务未配置]", "[AI 调用失败]"))

    def _extract_text(self, filepath: Path, suffix: str) -> str:
            if suffix in {".md", ".txt"}:
                return filepath.read_text(encoding="utf-8", errors="ignore")

            if suffix == ".docx":
                try:
                    from docx import Document
                except ImportError:
                    return "[缺少 python-docx，暂无法解析 docx 内容]"
                doc = Document(str(filepath))
                return "\n".join(p.text for p in doc.paragraphs if p.text).strip()

            if suffix == ".pdf":
                try:
                    from pypdf import PdfReader
                except ImportError:
                    return "[缺少 pypdf，暂无法解析 pdf 内容]"

                reader = PdfReader(str(filepath))
                text_parts = []
                for page in reader.pages:
                    text_parts.append(page.extract_text() or "")
                return "\n".join(text_parts).strip()

            return ""

